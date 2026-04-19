#!/usr/bin/env python3
"""Regenerate ORDER_BUDDY/Akram_performance_benchmark.docx (run from repo root or checkout-api)."""

from __future__ import annotations

import io
import json
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_para(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)


def add_bullets(doc: Document, items: list[str]) -> None:
    for t in items:
        doc.add_paragraph(t, style="List Bullet")


def _fmt_ms(x: float) -> str:
    return f"{x:.3f}"


def _load_json(path: Path) -> dict[str, Any] | None:
    if not path.is_file():
        return None
    try:
        return json.loads(path.read_text(encoding="utf-8"))
    except (json.JSONDecodeError, OSError):
        return None


def _rows_from_bench(data: dict[str, Any] | None, fallback: list[tuple[str, str]]) -> list[tuple[str, str]]:
    if not data:
        return fallback
    lat = data.get("latency_ms") or {}
    env = str(data.get("environment", ""))
    iters = data.get("iterations")
    warm = data.get("warmup")
    rows: list[tuple[str, str]] = [
        ("Environment", env or fallback[0][1]),
    ]
    if iters is not None and warm is not None:
        rows.append((f"Iterations (warmup {warm})", str(iters)))
    elif iters is not None:
        rows.append(("Iterations", str(iters)))
    for key, label in [
        ("min", "Min latency (ms)"),
        ("p50", "p50 (ms)"),
        ("mean", "Mean (ms)"),
        ("stdev", "Std dev (ms)"),
        ("p95", "p95 (ms)"),
        ("p99", "p99 (ms)"),
        ("max", "Max latency (ms)"),
    ]:
        if key in lat:
            rows.append((label, _fmt_ms(float(lat[key]))))
    tp = data.get("throughput_rps")
    if tp is not None:
        rows.append(("Approx. throughput (req/s, 1/mean)", f"{float(tp):.1f}"))
    co = data.get("clean_orders_before_timed_run")
    if co is not None:
        rows.append(("Clean orders before timed run", str(co)))
    return rows


def _latency_values(data: dict[str, Any]) -> dict[str, float]:
    lat = data.get("latency_ms") or {}
    out: dict[str, float] = {}
    for k in ("min", "p50", "mean", "p95", "p99", "max"):
        if k in lat:
            out[k] = float(lat[k])
    return out


def build_performance_plot_png(
    mem_json: dict[str, Any] | None,
    mongo_json: dict[str, Any] | None,
) -> io.BytesIO | None:
    """Return PNG bytes for a two-panel figure (latency bars + throughput), or None if nothing to plot."""
    if not mem_json and not mongo_json:
        return None
    try:
        import matplotlib

        matplotlib.use("Agg")
        import matplotlib.pyplot as plt
    except ImportError:
        return None

    labels = ["min", "p50", "mean", "p95", "p99", "max"]
    label_pretty = ["Min", "p50", "Mean", "p95", "p99", "Max"]
    series: list[tuple[str, dict[str, float], str]] = []
    if mem_json:
        series.append(("In-memory (fake Mongo)", _latency_values(mem_json), "#1f77b4"))
    if mongo_json:
        series.append(("Mongo (Motor + Docker)", _latency_values(mongo_json), "#ff7f0e"))

    fig, (ax_lat, ax_tp) = plt.subplots(2, 1, figsize=(8.2, 5.8), height_ratios=[2.2, 1])
    fig.suptitle("Checkout benchmark — latency and throughput (from JSON snapshots)", fontsize=11, y=0.98)

    x = range(len(labels))
    n = len(series)
    if n == 0:
        plt.close(fig)
        return None
    width = 0.75 / n
    for i, (name, vals, color) in enumerate(series):
        heights = [vals.get(k, 0.0) for k in labels]
        offset = (i - (n - 1) / 2) * width
        ax_lat.bar([xi + offset for xi in x], heights, width=width * 0.95, label=name, color=color)

    ax_lat.set_ylabel("Latency (ms)")
    ax_lat.set_xticks(list(x))
    ax_lat.set_xticklabels(label_pretty, rotation=0, fontsize=9)
    ax_lat.legend(loc="upper left", fontsize=8)
    ax_lat.grid(axis="y", linestyle="--", alpha=0.4)
    ax_lat.set_title("Request latency (summary statistics)", fontsize=10, pad=6)

    tp_names: list[str] = []
    tp_vals: list[float] = []
    tp_colors: list[str] = []
    if mem_json and mem_json.get("throughput_rps") is not None:
        tp_names.append("In-memory")
        tp_vals.append(float(mem_json["throughput_rps"]))
        tp_colors.append("#1f77b4")
    if mongo_json and mongo_json.get("throughput_rps") is not None:
        tp_names.append("Mongo")
        tp_vals.append(float(mongo_json["throughput_rps"]))
        tp_colors.append("#ff7f0e")

    if tp_names:
        ax_tp.bar(tp_names, tp_vals, color=tp_colors, width=0.55)
        ax_tp.set_ylabel("req/s")
        ax_tp.set_title("Approx. throughput (1 / mean latency)", fontsize=10, pad=6)
        ax_tp.grid(axis="y", linestyle="--", alpha=0.4)
    else:
        ax_tp.axis("off")

    fig.tight_layout(rect=[0, 0, 1, 0.96])
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=160, bbox_inches="tight")
    plt.close(fig)
    buf.seek(0)
    return buf


def add_performance_plots_section(
    doc: Document,
    mem_json: dict[str, Any] | None,
    mongo_json: dict[str, Any] | None,
) -> None:
    add_heading(doc, "Performance plots", level=2)
    png = build_performance_plot_png(mem_json, mongo_json)
    if png is None:
        add_para(
            doc,
            "Plots could not be generated: install matplotlib (pip install -r requirements-dev.txt) and "
            "ensure at least one of openapi/performance_results.json or openapi/performance_results_mongo.json exists.",
        )
        return
    add_para(
        doc,
        "Figure 1 — Latency summary (bars per statistic) and approximate throughput. Values are read from "
        "the same JSON files as the tables below; re-run the benchmark scripts then "
        "python scripts/build_akram_performance_doc.py to refresh.",
    )
    doc.add_picture(png, width=Inches(6.3))
    doc.add_paragraph()


def add_table_metrics(doc: Document, title: str, rows: list[tuple[str, str]]) -> None:
    add_heading(doc, title, level=2)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Metric"
    hdr[1].text = "Value"
    for k, v in rows:
        row = table.add_row().cells
        row[0].text = k
        row[1].text = v


def main() -> None:
    checkout_api = Path(__file__).resolve().parents[1]
    root = checkout_api.parents[1]  # ORDER_BUDDY
    out = root / "Akram_performance_benchmark.docx"

    mem_json = _load_json(checkout_api / "openapi/performance_results.json")
    mongo_json = _load_json(checkout_api / "openapi/performance_results_mongo.json")

    doc = Document()
    t = doc.add_heading("OrderBuddy: Order checkout performance analysis", 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_para(
        doc,
        "This document summarizes the NestJS-to-Python migration for order checkout, records "
        "reproducible latency measurements for the FastAPI service (orderbuddy/checkout-api), and "
        "includes an architectural audit. Benchmarks use bundled scripts; methodology must match "
        "when comparing against the legacy NestJS service.",
    )

    # ----- Task-1 -----
    add_heading(doc, "Task-1 — Migration approach, caveats, and tradeoffs", level=1)

    add_heading(doc, "How the migration was done", level=2)
    add_para(
        doc,
        "The migration targeted a single HTTP surface: POST /order-app/checkout from the NestJS "
        "OrderBuddy API. The Python implementation lives in orderbuddy/checkout-api as an async "
        "FastAPI application.",
    )
    add_bullets(
        doc,
        [
            "Stack: FastAPI, Pydantic v2 for request/response validation, Motor for async MongoDB "
            "access to the same collections (origins, menus, locations, orders) used by Nest.",
            "Service layer (CheckoutService) reproduces menu/origin resolution, availability checks, "
            "modifier-aware pricing, payment authorization, and order persistence.",
            "Outbound payment uses httpx with Tenacity retries and a circuit breaker; providers include "
            "mock (no network), generic HTTP JSON, and Emergepay-style integration with optional "
            "credentials loaded from the locations document.",
            "Contract parity is supported by exporting OpenAPI from both services and comparing checkout "
            "request/response schemas (see scripts/export_openapi.py, compare_openapi.py).",
            "Automated tests (pytest) cover the API, pricing, and OpenAPI comparison; CI can run the suite "
            "on changes under checkout-api.",
        ],
    )

    add_heading(doc, "Caveats and technical tradeoffs", level=2)
    add_bullets(
        doc,
        [
            "Scope: Only the checkout route was migrated in this slice; the broader Nest monolith and "
            "other endpoints are unchanged. Operating two stacks implies duplicate deployment, monitoring, "
            "and configuration until a full cutover.",
            "Validation semantics: Pydantic v2 and Nest/class-validator may differ slightly on edge cases; "
            "parity should be verified with real payloads and consumer expectations.",
            "Swagger coverage: If checkout is missing from the Nest OpenAPI export, DTOs/controllers must be "
            "annotated before schema comparison is meaningful.",
            "Benchmarks in this document: The in-memory benchmark replaces MongoDB with a fake in-process "
            "store; it measures Python + Pydantic + service logic, not database or network latency. The "
            "Mongo-backed script (Docker MongoDB + Motor) adds real read/write I/O but still uses "
            "FastAPI TestClient (in-process ASGI), not a full HTTP client load test against uvicorn.",
            "Production-like numbers require the same load generator (e.g. hey, k6), payload, concurrency, "
            "MongoDB topology, and payment mode on both Nest and FastAPI.",
            "Payment mocking: Benchmarks default to PAYMENT_GATEWAY_MOCK=true to isolate application work; "
            "enabling real PSP calls introduces network variance and compliance constraints.",
        ],
    )

    # ----- Task-2 -----
    add_heading(doc, "Task-2 — Architectural audit", level=1)

    add_heading(doc, "System fit and data model", level=2)
    add_para(
        doc,
        "The FastAPI service reuses the existing MongoDB schema for checkout: reads from origins, menus, "
        "and locations; writes orders with the same field shapes expected downstream. No separate "
        "microservice database was introduced for this slice, reducing synchronization risk with the "
        "Nest deployment.",
    )

    add_heading(doc, "Resilience and integrations", level=2)
    add_bullets(
        doc,
        [
            "Payment calls are isolated behind CheckoutPaymentClient with configurable timeouts, retries, "
            "and a circuit breaker to avoid cascading failures when a PSP is slow or unavailable.",
            "Async I/O (Motor, httpx) aligns with concurrent request handling under uvicorn or similar ASGI "
            "servers.",
        ],
    )

    add_heading(doc, "Security and configuration", level=2)
    add_bullets(
        doc,
        [
            "Secrets and endpoints are environment-driven (e.g. MONGODB_URI, Emergepay URLs/tokens); "
            ".env must not be committed.",
            "Emergepay credentials can be resolved from the location document to match Nest behavior "
            "when configured.",
        ],
    )

    add_heading(doc, "Observability and operations (gaps and recommendations)", level=2)
    add_bullets(
        doc,
        [
            "Recommend structured logging with correlation IDs across checkout and payment calls for "
            "production troubleshooting.",
            "Recommend metrics (latency histograms, error rates, circuit-breaker state) and optional "
            "distributed tracing for parity with existing Nest observability.",
            "Load and soak testing should be repeated after any payment-provider or schema change.",
        ],
    )

    add_heading(doc, "Maintainability", level=2)
    add_para(
        doc,
        "Code is split into schemas, services, payment gateway adapter, and configuration. Shared "
        "benchmark fixtures (app/benchmark_fixtures.py) keep tests, seeds, and scripts aligned.",
    )

    # ----- Methodology -----
    add_heading(doc, "Performance methodology", level=1)
    add_para(doc, "Run from orderbuddy/checkout-api with a Python 3.12+ venv and dependencies installed.")
    add_para(doc, "A — In-process benchmark (fake MongoDB, mock payment):")
    add_para(
        doc,
        "python scripts/benchmark_checkout.py --iterations 1500 --warmup 100 "
        "--json-out openapi/performance_results.json",
    )
    add_para(doc, "B — MongoDB-backed benchmark (Docker MongoDB, Motor, mock payment):")
    add_para(
        doc,
        "docker compose up -d && export MONGODB_URI=mongodb://127.0.0.1:27017 "
        "MONGODB_DB=orderbuddy_bench && python scripts/seed_mongo_benchmark.py && "
        "python scripts/benchmark_checkout_mongo.py --iterations 1500 --warmup 100 "
        "--json-out openapi/performance_results_mongo.json",
    )

    # ----- Metrics (representative snapshots; regenerate JSON on your machine) -----
    add_heading(doc, "Captured metrics (representative snapshots)", level=1)
    add_para(
        doc,
        "Figures below are examples from scripted runs; regenerate openapi/performance_results*.json on "
        "your hardware for reporting. In-memory metrics isolate CPU/Python work; Mongo metrics include "
        "database I/O.",
    )

    add_performance_plots_section(doc, mem_json, mongo_json)

    add_table_metrics(
        doc,
        "A — In-memory benchmark (see openapi/performance_results.json)",
        _rows_from_bench(
            mem_json,
            [
                ("Environment", "FastAPI TestClient, in-memory fake DB, mock payment"),
                ("Min latency (ms)", "—"),
                ("p50 (ms)", "—"),
                ("Mean (ms)", "—"),
                ("Std dev (ms)", "—"),
                ("p95 (ms)", "—"),
                ("p99 (ms)", "—"),
                ("Max latency (ms)", "—"),
                ("Approx. throughput (req/s, 1/mean)", "—"),
            ],
        ),
    )
    doc.add_paragraph()

    add_table_metrics(
        doc,
        "B — Mongo-backed benchmark (see openapi/performance_results_mongo.json)",
        _rows_from_bench(
            mongo_json,
            [
                ("Environment", "FastAPI TestClient + Motor, MongoDB (e.g. Docker), mock payment"),
                ("Mean latency (ms)", "—"),
                ("Approx. throughput (req/s)", "—"),
            ],
        ),
    )
    add_para(
        doc,
        "Regenerate the JSON files with the benchmark scripts on your machine, then re-run: "
        "python scripts/build_akram_performance_doc.py. Mongo metrics vary with disk, tmpfs vs volume, "
        "and host load.",
    )

    add_heading(doc, "Interpretation", level=1)
    add_para(
        doc,
        "In-process benchmarks are best for regression detection on the Python checkout path. They do not "
        "replace a fair Nest vs FastAPI comparison unless both sides use the same load tool, payload, "
        "MongoDB configuration, and payment mode.",
    )

    add_heading(doc, "Nest comparison", level=1)
    add_para(
        doc,
        "For stakeholder comparison, capture comparable runs (e.g. median or p95 latency) using the same "
        "methodology on both services, or record a short walkthrough with identical payloads and backend data.",
    )

    add_para(doc, "— End of document —")

    doc.save(out)
    print(f"Wrote {out}")


if __name__ == "__main__":
    main()
